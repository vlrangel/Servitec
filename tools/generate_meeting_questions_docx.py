from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path("docs/Guion-Reunion-Cliente-Partner-BC-NewSecuryTechnics.docx")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: object, bold: bool = False, white: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    if white:
        run.font.color.rgb = RGBColor(255, 255, 255)


def add_table(document: Document, headers: list[str], rows: list[tuple[object, ...]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(header_cells[index], header, bold=True, white=True)
        set_cell_shading(header_cells[index], "1F4E79")
        header_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def main() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(12)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Guion de reunion: cliente + partner Business Central")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("NewSecuryTechnics - Validacion del presupuesto Python + BC")
    run.font.size = Pt(12)

    date_p = document.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(f"Fecha: {date.today().strftime('%d/%m/%Y')}")

    document.add_heading("1. Objetivo de la reunion", level=1)
    document.add_paragraph(
        "Validar que el presupuesto comercial de 36.000 EUR + IVA es defendible porque "
        "Business Central actua como backoffice y unica fuente de datos, y Python solo "
        "implementa la capa web. Para cerrarlo hay que confirmar que logica ya existe en AL "
        "del partner, que APIs/OData hay publicadas, y si se puede usar su extension como dependencia."
    )
    document.add_paragraph(
        "Si la logica critica vive solo en la intranet PHP, el coste no es de web: es desarrollo AL en BC y debe tratarse como ampliacion."
    )

    document.add_heading("2. Mensaje de apertura (30 segundos)", level=1)
    document.add_paragraph(
        "El presupuesto asume que Business Central es el backoffice y la unica fuente de datos. "
        "Python no guarda datos de negocio. Para cerrarlo necesitamos saber que logica ya esta "
        "en AL del partner, que APIs/OData existen, y si podemos usar su extension como dependencia."
    )

    document.add_heading("3. Preguntas criticas sobre la extension del partner", level=1)
    document.add_paragraph("Prioridad maxima. Conseguir respuesta clara Si / No / Condicionado.")
    add_numbered(
        document,
        [
            "Podemos declarar su extension como dependencia en una extension nuestra (app.json -> dependencies)?",
            "Si si: nos dan App ID, publisher, nombre y version minima?",
            "Si si: es publica/compartible o solo esta en su entorno?",
            "Si si: nos entregan el .app, fuente y/o documentacion de objetos?",
            "Si si: que objetos son publicos (tables, pages, codeunits, enums) y cuales son internos?",
            "Si no: publican ellos APIs/OData/SOAP para consumir desde Python sin depender de su app?",
            "Si no: quien desarrolla y mantiene los endpoints nuevos?",
            "Licencia y propiedad: podemos modificar/extender, o solo consumir?",
            "Compatibilidad: version exacta de BC (online/on-prem), runtime y numero de companias?",
            "Hay sandbox/preproduccion con la misma extension instalada?",
            "Permisos: podemos publicar web services / API pages / codeunits, o solo el partner?",
        ],
    )

    document.add_heading("4. Pregunta clave de arquitectura", level=1)
    document.add_paragraph(
        "Para cada modulo del presupuesto, pedir una de estas tres respuestas:"
    )
    add_table(
        document,
        ["Respuesta", "Implicacion para el presupuesto"],
        [
            ("Ya esta en AL (tablas, paginas, codeunits, triggers)", "Python solo UI + llamadas. Presupuesto OK."),
            ("Esta en PHP y BC solo guarda maestros", "Hay que trasladar logica a AL. Riesgo / ampliacion."),
            ("Mitad y mitad", "Pedir mapa: que valida BC y que hace la intranet."),
        ],
    )
    document.add_paragraph("")
    document.add_paragraph(
        "Pregunta literal a formular:"
    )
    document.add_paragraph(
        "Que parte de la estructura y de la logica de negocio de cada punto esta desarrollada "
        "en AL de Business Central, y que parte vive hoy solo en la intranet PHP?"
    )

    document.add_heading("5. Preguntas por modulo (tabla del presupuesto)", level=1)
    document.add_paragraph(
        "Recorrer M1 a M13. Por cada modulo, preguntar:"
    )
    add_numbered(
        document,
        [
            "Existe tabla/entidad en BC?",
            "Existe pagina o API/OData publicada?",
            "La logica (estados, asignar, terminar, registrar, stock, PDF...) esta en codeunit AL o en PHP?",
            "Se puede escribir desde fuera (POST/PATCH/accion) o solo leer?",
            "Quien mantiene cambios: partner, nosotros, o ambos?",
        ],
    )

    document.add_paragraph("")
    document.add_paragraph("Checklist por modulo (rellenar en la reunion):")
    add_table(
        document,
        ["Modulo", "Nombre", "AL / PHP / Mixto", "API/OData existe", "Escritura externa", "Notas"],
        [
            ("M1", "Acceso, usuarios y seguridad", "", "", "", ""),
            ("M2", "Dashboard e intranet principal", "", "", "", ""),
            ("M3", "Incidencias SAT", "", "", "", ""),
            ("M4", "Mantenimientos", "", "", "", ""),
            ("M5", "Obras, proyectos y partes de proyecto", "", "", "", ""),
            ("M6", "Partes de trabajo de tecnico", "", "", "", ""),
            ("M7", "Asignacion de ordenes pendientes y rutas", "", "", "", ""),
            ("M8", "Inventario, almacenes y centros", "", "", "", ""),
            ("M9", "Pedidos de transferencia", "", "", "", ""),
            ("M10", "Antihurtos y checklist tecnico", "", "", "", ""),
            ("M11", "Presupuestos y albaranes", "", "", "", ""),
            ("M12", "Documentacion, planos y archivos", "", "", "", ""),
            ("M13", "Costes, gastos e informes", "", "", "", ""),
        ],
    )

    document.add_paragraph("")
    document.add_paragraph("Modulos donde mas duele si no estan en AL:")
    add_bullets(
        document,
        [
            "M3 Incidencias / M4 Mantenimientos / M5 Obras: estados, verificacion, descarte, registrar albaran.",
            "M6 Partes de trabajo: materiales, firma, tiempos, PDF.",
            "M7 Asignacion y rutas: asignacion tecnico + ruta.",
            "M8 Inventario / M9 Transferencias: stock y movimientos.",
            "M10 Antihurtos/checklist: muy probablemente custom.",
            "M11 Presupuestos/albaranes: documentos de venta.",
            "M12 Documentos/planos: adjuntos BC o ficheros en servidor PHP?",
        ],
    )

    document.add_heading("6. Integracion tecnica", level=1)
    add_numbered(
        document,
        [
            "Que hay publicado hoy: API pages, OData, SOAP, codeunits?",
            "Existe algo tipo incWebServicesAPI.php en la intranet que ya llame a BC? A que endpoints?",
            "Autenticacion: OAuth2 / Entra ID, usuario web service, NavUserPassword, NTLM...?",
            "Hay usuario de servicio para la nueva web?",
            "Limites: paginacion, timeouts, volumen de listados (incidencias, inventario).",
            "Adjuntos: Blob Storage BC, media, o filesystem de la intranet?",
            "PDFs/albaranes: los genera BC o PHP?",
            "Multiempresa: empresaNavision implica varias companias? Como se selecciona?",
        ],
    )

    document.add_heading("7. Gobernanza y roles", level=1)
    add_numbered(
        document,
        [
            "Quien hace el desarrollo AL nuevo: partner, nosotros, o mixto?",
            "Si el partner hace AL y nosotros Python: plazos y coste del partner estan fuera de los 36.000 EUR?",
            "Accesos: sandbox BC, usuario tecnico, AppSource/private app, Azure DevOps/Git del partner.",
            "Hay documentacion funcional/tecnica de su extension?",
            "Pueden ensenar en vivo: Web Services page + objetos de ordenes/partes/inventario?",
        ],
    )

    document.add_heading("8. Preguntas para defender los 36.000 EUR", level=1)
    document.add_paragraph("Dejar claros estos supuestos:")
    add_numbered(
        document,
        [
            "No hay BD propia de negocio en Python; si hace falta, se abre ampliacion.",
            "No se redisena el proceso; se replica lo existente.",
            "El bloque Backoffice BC del presupuesto cubre publicar/ajustar APIs, no rehacer todo el SAT en AL desde cero.",
            "Si la logica critica esta solo en PHP, el presupuesto de web no incluye reescribirla entera en AL.",
        ],
    )
    document.add_paragraph("")
    document.add_paragraph("Frase util:")
    document.add_paragraph(
        "Los 36.000 EUR asumen reutilizar estructura y logica ya existente en BC/AL o APIs del partner. "
        "Si hay que desarrollar en AL tablas, codeunits y procesos que hoy solo existen en PHP, "
        "eso es un alcance AL aparte."
    )

    document.add_heading("9. Preguntas trampa al partner", level=1)
    add_numbered(
        document,
        [
            "Si manana cae la intranet PHP, que procesos se pueden seguir operando solo desde BC?",
            "Que codeunits se ejecutan al terminar un parte / registrar una orden?",
            "El stock de centro/delegacion es stock estandar BC o tablas custom vuestras?",
            "Los numeros OT26-... / PV26-... / OFP-... se numeran en BC o en PHP?",
            "Podeis publicar una API de lectura y otra de escritura por entidad critica en 2-3 semanas?",
        ],
    )

    document.add_heading("10. Checklist de salida (debe quedar escrito)", level=1)
    add_bullets(
        document,
        [
            "Extension partner usable como dependency (si/no + App ID).",
            "Acceso a .app / objetos / documentacion.",
            "Lista de web services/APIs existentes.",
            "Por modulo M1-M13: AL / PHP / mixto.",
            "Quien desarrolla AL nuevo.",
            "Entorno sandbox + usuario integracion.",
            "Metodo de autenticacion.",
            "Adjuntos y PDFs: donde viven.",
            "Confirmacion de que el alcance es equivalente, no mejora de proceso.",
        ],
    )

    document.add_heading("11. Orden recomendado de la reunion (45-60 min)", level=1)
    add_table(
        document,
        ["Minutos", "Bloque", "Que conseguir"],
        [
            ("5", "Enfoque Python sin BD + BC backoffice", "Alinear expectativa del presupuesto."),
            ("15", "Dependency + APIs existentes", "Saber si reutilizamos extension/partner."),
            ("25", "Recorrido M1-M13: AL o PHP", "Rellenar tabla de modulos."),
            ("10", "Roles, plazos y riesgos al precio", "Separar coste web vs coste AL."),
        ],
    )

    document.add_heading("12. Notas de la reunion", level=1)
    document.add_paragraph("Asistentes:")
    document.add_paragraph("")
    document.add_paragraph("Decisiones / acuerdos:")
    document.add_paragraph("")
    document.add_paragraph("")
    document.add_paragraph("Riesgos detectados:")
    document.add_paragraph("")
    document.add_paragraph("")
    document.add_paragraph("Proximos pasos:")
    document.add_paragraph("")
    document.add_paragraph("")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
