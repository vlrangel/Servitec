from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path("docs/Presupuesto-Migracion-Intranet-Python-BC-NewSecuryTechnics.docx")
RATE = 600


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: object, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    if bold:
        run.font.color.rgb = RGBColor(255, 255, 255)


def money(value: int) -> str:
    return f"{value:,.0f} EUR".replace(",", ".")


def add_table(document: Document, headers: list[str], rows: list[tuple[object, ...]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(header_cells[index], header, bold=True)
        set_cell_shading(header_cells[index], "1F4E79")
        header_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


modules = [
    ("M1", "Acceso, usuarios y seguridad", "Login, perfiles, cambio de clave, permisos por rol, sesiones, HTTPS y auditoria basica.", "Python + BC usuarios/permisos"),
    ("M2", "Dashboard e intranet principal", "Pantalla principal, menu, notificaciones, accesos por perfil y resumen operativo.", "Python consumiendo datos BC"),
    ("M3", "Incidencias SAT", "Listado, filtros por fechas/cliente/centro/tecnico/estado, asignacion, verificacion, descarte, finalizacion, documentos y albaran SAT.", "BC ordenes SAT + APIs/OData"),
    ("M4", "Mantenimientos", "Gestion equivalente a incidencias para mantenimientos: busqueda, asignacion, cierre, resoluciones, adjuntos y registro.", "BC ordenes/mantenimientos"),
    ("M5", "Obras, proyectos y partes de proyecto", "Gestion de ordenes de proyecto, proyectos, partes asociados, albaranes de proyecto, documentos y consulta historica.", "BC proyectos/obras + APIs"),
    ("M6", "Partes de trabajo de tecnico", "Alta/edicion de parte, materiales usados/recogidos/reutilizados, firma, observaciones, tiempos, documentos y PDF.", "Python UI + escritura directa en BC via servicios"),
    ("M7", "Asignacion de ordenes pendientes y rutas", "Ordenes pendientes, asignacion a tecnicos, ruta diaria por tecnico/delegacion/gerente/cliente y ocupacion de tecnicos.", "BC agenda/recursos/rutas"),
    ("M8", "Inventario, almacenes y centros", "Inventario general, inventario por centro, inventario por delegacion, materiales, movimientos y busquedas.", "BC items/almacenes/stock"),
    ("M9", "Pedidos de transferencia", "Origen/destino, lineas de productos, envio, recepcion, cancelacion y estados del pedido.", "BC transfer orders o desarrollo equivalente"),
    ("M10", "Antihurtos y checklist tecnico", "Inventario de antenas/desactivadores antihurto, checklist NST 2026, observaciones y estado de instalacion.", "BC tablas/backoffice especifico"),
    ("M11", "Presupuestos y albaranes", "Listado, filtros, copia/consulta de presupuestos, generacion de albaranes SAT/proyecto y PDFs.", "BC ventas/proyectos + generacion documentos"),
    ("M12", "Documentacion, planos y archivos", "Documentacion de centros, planos, subida/consulta de archivos, lectura de documentos y permisos.", "BC adjuntos o repositorio documental enlazado desde BC"),
    ("M13", "Costes, gastos e informes", "Resumen de costes, cuentas de gastos, filtros por gerente/empresa/fechas y exportaciones.", "Consultas BC/API pages"),
]

budget_items = [
    ("Analisis funcional y cierre de alcance", "Revision asistida de pantallas, priorizacion de modulos y matriz pantalla -> entidad BC.", "2.400 EUR"),
    ("Base tecnica Python sin base de datos propia", "Proyecto base, seguridad, usuarios, permisos, plantillas, servicios BC y despliegue inicial.", "4.800 EUR"),
    ("Backoffice Business Central", "Publicacion/creacion de APIs, paginas OData, codeunits y ajustes necesarios para usar BC como unico origen de datos.", "7.200 EUR"),
    ("Migracion funcional de modulos", "Desarrollo de pantallas Python y flujos contra BC para los modulos M1-M13, apoyado por IA para acelerar maquetacion, formularios y codigo repetitivo.", "16.800 EUR"),
    ("Pruebas, puesta en marcha y formacion", "Pruebas funcionales, correcciones, despliegue en produccion, documentacion y formacion breve.", "4.800 EUR"),
]

total_amount = 36_000


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def main() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "Calibri"
    styles["Title"].font.size = Pt(20)
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(12)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Presupuesto migracion intranet a Python + Business Central")
    run.bold = True
    run.font.size = Pt(20)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("NewSecuryTechnics - Propuesta economica y plan de ejecucion")
    run.font.size = Pt(12)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(f"Fecha: {date.today().strftime('%d/%m/%Y')}")

    document.add_heading("1. Resumen ejecutivo", level=1)
    document.add_paragraph(
        "Se propone sustituir la intranet PHP actual publicada en https://bc-nav.ns-technic.com por una aplicacion web moderna en Python, manteniendo Business Central/NAV como unico backoffice y origen de datos de negocio."
    )
    document.add_paragraph(
        "La nueva aplicacion Python no tendra base de datos de negocio propia. Las pantallas consultaran y modificaran informacion mediante APIs, OData o servicios publicados en Business Central. Python actuara como capa web, seguridad, experiencia de usuario, validacion de formularios, generacion de vistas y orquestacion de llamadas."
    )
    document.add_paragraph(
        f"Presupuesto comercial reducido: {money(total_amount)} + IVA. Plazo de calendario estimado: 10 a 12 semanas, considerando desarrollo asistido por IA, reutilizacion de patrones y trabajo en paralelo entre Python, backoffice Business Central y pruebas."
    )

    document.add_heading("2. Informacion usada para la estimacion", level=1)
    add_bullets(
        document,
        [
            "Auditoria autenticada de solo lectura sobre la intranet actual, con 80 paginas revisadas y sin envio de formularios funcionales.",
            "Pantallas detectadas: incidencias, mantenimientos, obras/proyectos, partes de trabajo, inventario, pedidos de transferencia, antihurtos, rutas, ocupacion, presupuestos, costes, documentacion, planos, ordenes pendientes y gestion de clave.",
            "Formularios detectados con filtros, adjuntos, resoluciones, asignaciones, partes de tecnico, checklist, materiales, transferencias y generacion/consulta de documentos.",
            "La propuesta se ha ajustado a un enfoque de coste reducido usando IA para acelerar programacion, maquetacion, formularios y codigo repetitivo.",
            "La estimacion queda condicionada a confirmar en Business Central las entidades disponibles, permisos, version, autenticacion y endpoints existentes.",
        ],
    )

    document.add_heading("3. Enfoque tecnico", level=1)
    add_bullets(
        document,
        [
            "Frontend/backend web en Python, recomendado Django por usuarios, permisos, formularios, plantillas, panel administrativo tecnico y seguridad.",
            "Sin base de datos propia para datos de negocio: clientes, centros, ordenes, partes, inventario, presupuestos, costes y documentos residiran en Business Central/NAV.",
            "Se podra usar almacenamiento tecnico minimo no funcional, si fuera imprescindible, solo para cache temporal, sesiones, logs tecnicos o colas; no como maestro de datos.",
            "Backoffice en Business Central: publicacion o desarrollo de APIs, paginas, consultas, codeunits y pantallas internas para mantener datos maestros y procesos.",
            "Integracion preferente via APIs REST/OData. SOAP solo si la version de NAV/BC lo exige.",
            "Autenticacion de integracion mediante OAuth2 si el entorno lo permite; alternativa con usuario de servicio/web service para NAV/BC on-premise.",
        ],
    )

    document.add_heading("4. Modulos a migrar a Python con backoffice en Business Central", level=1)
    module_rows = [(code, name, scope, bc) for code, name, scope, bc in modules]
    add_table(document, ["Modulo", "Nombre", "Alcance Python", "Backoffice / datos en BC"], module_rows)

    document.add_heading("5. Desglose economico", level=1)
    budget_rows = [(name, desc, amount) for name, desc, amount in budget_items]
    add_table(document, ["Concepto", "Descripcion", "Importe"], budget_rows)
    document.add_paragraph(f"Total estimado: {money(total_amount)} + IVA.")
    document.add_paragraph(
        "Este importe se plantea como presupuesto comercial reducido, apoyado en programacion asistida por IA y en que Business Central actue como unico backoffice y fuente de datos."
    )

    document.add_heading("6. Plan de ejecucion", level=1)
    add_table(
        document,
        ["Fase", "Nombre", "Plazo", "Entregable"],
        [
            ("Fase 1", "Analisis y diseno funcional/tecnico", "Semanas 1-2", "Mapa de pantallas, matriz de entidades BC, alcance cerrado."),
            ("Fase 2", "Base tecnica Python + integracion BC", "Semanas 2-3", "Proyecto base, autenticacion, permisos, servicios BC y despliegue preproduccion."),
            ("Fase 3", "Backoffice Business Central", "Semanas 3-5", "APIs/OData/codeunits/paginas BC para soportar los modulos."),
            ("Fase 4", "Migracion de modulos principales", "Semanas 4-9", "Incidencias, mantenimientos, obras/proyectos, partes, inventario, rutas y presupuestos."),
            ("Fase 5", "Pruebas, UAT y puesta en marcha", "Semanas 9-12", "Validacion con usuarios, correcciones, formacion y produccion."),
        ],
    )

    document.add_heading("7. Alcance incluido", level=1)
    add_bullets(
        document,
        [
            "Migracion funcional de los modulos detectados en la auditoria autenticada.",
            "Desarrollo de pantallas Python responsive para uso de oficina y tecnicos.",
            "Lectura y escritura contra Business Central/NAV mediante servicios publicados o desarrollados.",
            "Backoffice en Business Central para datos maestros y configuraciones necesarias.",
            "Gestion de usuarios, permisos, errores, trazabilidad tecnica y configuracion segura.",
            "Generacion o consulta de documentos principales: partes, albaranes, presupuestos y archivos asociados, segun disponibilidad de datos en BC.",
            "Entorno de preproduccion, despliegue en produccion, documentacion y formacion breve.",
        ],
    )

    document.add_heading("8. Exclusiones y condicionantes", level=1)
    add_bullets(
        document,
        [
            "No incluye licencias de Microsoft, usuarios Business Central, hosting, dominios ni certificados.",
            "No incluye redisenar procesos de negocio no existentes en la intranet actual.",
            "El precio reducido exige mantener el alcance funcional equivalente a la intranet actual y evitar cambios de proceso durante el desarrollo.",
            "No incluye migracion historica desde una base de datos propia de la intranet; si existe historico externo y debe migrarse a BC, se presupuestara aparte.",
            "No incluye integraciones con terceros distintas de Business Central/NAV salvo que se documenten en el analisis.",
            "No incluye soporte 24x7 ni mantenimiento evolutivo posterior.",
            "El importe requiere confirmar version de BC/NAV, autenticacion, endpoints, companias, permisos y posibilidad de crear/publicar objetos en BC.",
            "Si Business Central no contiene actualmente alguna entidad necesaria, habra que crearla en BC como parte del backoffice o presupuestarla como ampliacion.",
        ],
    )

    document.add_heading("9. Costes recurrentes recomendados", level=1)
    add_table(
        document,
        ["Concepto", "Descripcion", "Estimacion"],
        [
            ("Hosting / servidor gestionado", "Servidor aplicacion Python, HTTPS, copias y monitorizacion basica.", "80 a 200 EUR/mes"),
            ("Mantenimiento correctivo", "Correcciones, ajustes menores y soporte funcional.", "600 a 1.000 EUR/mes"),
            ("Bolsa evolutiva opcional", "Nuevas pantallas, cambios de flujo o mejoras tras puesta en marcha.", "Segun consumo"),
        ],
    )

    document.add_heading("10. Validez de la propuesta", level=1)
    document.add_paragraph(
        "Propuesta valida durante 30 dias. El presupuesto se considera estimacion cerrable tras una sesion de validacion funcional con usuarios clave y una revision tecnica del entorno Business Central/NAV."
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)
    print(f"Importe: {total_amount}")


if __name__ == "__main__":
    main()
